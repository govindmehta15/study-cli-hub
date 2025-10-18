# file_viewer.py
import os, csv, sys, getch
from rich.console import Console
from rich.panel import Panel
from rich.text import Text
from rich.prompt import Prompt
from datetime import datetime
from error_handler import handle_error

console = Console()

try:
    import PyPDF2
except:
    PyPDF2 = None

try:
    import docx
except:
    docx = None

def clear_screen():
    """Clear the terminal screen"""
    os.system("cls" if os.name == "nt" else "clear")

def get_key():
    """Cross-platform key input"""
    try:
        import msvcrt
        if msvcrt.kbhit():
            key = msvcrt.getch()
            if key == b'\xe0':  # Special key
                key = msvcrt.getch()
                if key == b'H': return 'up'
                elif key == b'P': return 'down'
                elif key == b'I': return 'page_up'
                elif key == b'Q': return 'page_down'
            elif key == b' ': return 'space'
            elif key == b'q': return 'quit'
            elif key == b'/': return 'search'
            elif key == b'h': return 'help'
            elif key == b'\r': return 'enter'
    except ImportError:
        # Unix/Linux/Mac
        import termios, tty
        fd = sys.stdin.fileno()
        old_settings = termios.tcgetattr(fd)
        try:
            tty.setraw(sys.stdin.fileno())
            ch = sys.stdin.read(1)
            if ch == '\x1b':  # ESC sequence
                ch = sys.stdin.read(2)
                if ch == '[A': return 'up'
                elif ch == '[B': return 'down'
                elif ch == '[5~': return 'page_up'
                elif ch == '[6~': return 'page_down'
            elif ch == ' ': return 'space'
            elif ch == 'q': return 'quit'
            elif ch == '/': return 'search'
            elif ch == 'h': return 'help'
            elif ch == '\r': return 'enter'
        finally:
            termios.tcsetattr(fd, termios.TCSADRAIN, old_settings)
    return None

def interactive_text_reader(path, filename):
    """Interactive text file reader with navigation and highlighting"""
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            lines = f.readlines()
        
        highlighted = set()
        search_results = []
        search_index = 0
        current_line = 0
        lines_per_page = 20
        
        while True:
            console.clear()
            
            # Header
            header = f"[bold cyan]{filename}[/bold cyan] | Line {current_line + 1}/{len(lines)}"
            if search_results:
                header += f" | Search: {len(search_results)} results"
            console.print(Panel(header, expand=False))
            
            # Help
            console.print("[dim]↑↓: Scroll | PgUp/PgDn: Page | Space: Highlight | /: Search | h: Help | q: Quit[/dim]")
            console.print()
            
            # Display lines
            start_line = max(0, current_line - lines_per_page // 2)
            end_line = min(len(lines), start_line + lines_per_page)
            
            for i in range(start_line, end_line):
                line_no = f"{i+1:>4}"
                line_text = lines[i].rstrip()
                
                # Highlighting
                if i in highlighted:
                    console.print(f"[bold black on yellow]{line_no} {line_text}[/bold black on yellow]")
                elif i == current_line:
                    console.print(f"[bold white on blue]{line_no} {line_text}[/bold white on blue]")
                else:
                    console.print(f"[green]{line_no}[/green] {line_text}")
            
            # Footer
            if current_line >= len(lines) - 1:
                console.print("\n[bold yellow]End of file[/bold yellow]")
            
            # Get user input
            key = get_key()
            if not key:
                continue
                
            if key == 'up':
                current_line = max(0, current_line - 1)
            elif key == 'down':
                current_line = min(len(lines) - 1, current_line + 1)
            elif key == 'page_up':
                current_line = max(0, current_line - lines_per_page)
            elif key == 'page_down':
                current_line = min(len(lines) - 1, current_line + lines_per_page)
            elif key == 'space':
                # Toggle highlight for current line
                if current_line in highlighted:
                    highlighted.remove(current_line)
                else:
                    highlighted.add(current_line)
            elif key == 'search':
                search_term = Prompt.ask("[yellow]Enter search term[/yellow]").strip()
                if search_term:
                    search_results = []
                    for i, line in enumerate(lines):
                        if search_term.lower() in line.lower():
                            search_results.append(i)
                    if search_results:
                        current_line = search_results[0]
                        console.print(f"[green]Found {len(search_results)} results[/green]")
                    else:
                        console.print("[red]No results found[/red]")
                        input("Press Enter to continue...")
            elif key == 'help':
                show_reader_help()
            elif key == 'quit':
                break
                
    except Exception as e:
        handle_error(e)

def show_reader_help():
    """Show help for the interactive reader"""
    console.clear()
    console.print(Panel("[bold cyan]Interactive Reader Help[/bold cyan]", expand=False))
    console.print()
    console.print("[green]Navigation:[/green]")
    console.print("  ↑ / ↓     - Scroll line by line")
    console.print("  PgUp/PgDn - Jump multiple lines")
    console.print()
    console.print("[yellow]Features:[/yellow]")
    console.print("  SPACE     - Highlight/unhighlight current line")
    console.print("  /         - Search for text in file")
    console.print("  h         - Show this help")
    console.print("  q         - Quit reader")
    console.print()
    input("Press Enter to continue...")

def edit_file_with_reason(user_folder, subject, filename):
    """Edit a file with reason tracking"""
    path = os.path.join("subjects", user_folder, subject, filename) if user_folder else os.path.join("subjects", subject, filename)
    
    if not os.path.exists(path):
        console.print(f"[red]File '{filename}' not found[/red]")
        return
    
    try:
        # Get edit reason
        reason = Prompt.ask("[yellow]Enter reason for editing this file[/yellow]").strip()
        if not reason:
            reason = "No reason provided"
        
        # Read current content
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        
        # Show current content
        console.print(Panel(f"[bold cyan]Editing: {filename}[/bold cyan]", expand=False))
        console.print(f"[dim]Reason: {reason}[/dim]")
        console.print()
        console.print("[yellow]Current content:[/yellow]")
        console.print(content)
        console.print()
        
        # Get new content
        console.print("[yellow]Enter new content (press Ctrl+D when done):[/yellow]")
        new_content = ""
        try:
            while True:
                line = input()
                new_content += line + "\n"
        except EOFError:
            pass
        
        # Save with reason tracking
        backup_path = path + f".backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        with open(backup_path, "w", encoding="utf-8") as f:
            f.write(content)
        
        with open(path, "w", encoding="utf-8") as f:
            f.write(new_content)
        
        # Log the edit
        log_path = os.path.join("subjects", user_folder, subject, "edit_log.txt") if user_folder else os.path.join("subjects", subject, "edit_log.txt")
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(f"[{datetime.now()}] Edited {filename} - Reason: {reason}\n")
            f.write(f"Backup saved as: {os.path.basename(backup_path)}\n\n")
        
        console.print(f"[green]✅ File edited successfully![/green]")
        console.print(f"[dim]Backup saved as: {os.path.basename(backup_path)}[/dim]")
        
    except Exception as e:
        handle_error(e)

def view_binary_file(path, filename, ext):
    """View binary files as hex dump"""
    try:
        console.print(f"[yellow]Binary file detected. Showing hex dump:[/yellow]")
        console.print()
        
        with open(path, "rb") as f:
            chunk_size = 16
            offset = 0
            
            for chunk in iter(lambda: f.read(chunk_size), b""):
                # Create hex representation
                hex_str = " ".join(f"{b:02x}" for b in chunk)
                hex_str = hex_str.ljust(47)  # Pad to align ASCII
                
                # Create ASCII representation
                ascii_str = "".join(chr(b) if 32 <= b <= 126 else "." for b in chunk)
                
                console.print(f"{offset:08x}: {hex_str} |{ascii_str}|")
                offset += len(chunk)
                
                # Limit output for very large files
                if offset > 1024:  # Show first 1KB
                    console.print(f"\n[yellow]... (showing first 1KB of {os.path.getsize(path)} bytes)[/yellow]")
                    break
        
        input("\nPress Enter to continue...")
        
    except Exception as e:
        console.print(f"[red]Error reading binary file: {e}[/red]")
        input("Press Enter to continue...")

def view_image_info(path, filename, ext):
    """Show image file information"""
    try:
        file_size = os.path.getsize(path)
        console.print(f"[cyan]Image file: {filename}[/cyan]")
        console.print(f"[cyan]Format: {ext.upper()}[/cyan]")
        console.print(f"[cyan]Size: {file_size:,} bytes[/cyan]")
        console.print()
        console.print("[yellow]Image files cannot be displayed in CLI.[/yellow]")
        console.print("[yellow]Use an image viewer or convert to text format.[/yellow]")
        input("\nPress Enter to continue...")
        
    except Exception as e:
        console.print(f"[red]Error reading image file: {e}[/red]")
        input("Press Enter to continue...")

def view_archive_info(path, filename, ext):
    """Show archive file information"""
    try:
        file_size = os.path.getsize(path)
        console.print(f"[cyan]Archive file: {filename}[/cyan]")
        console.print(f"[cyan]Format: {ext.upper()}[/cyan]")
        console.print(f"[cyan]Size: {file_size:,} bytes[/cyan]")
        console.print()
        console.print("[yellow]Archive files cannot be extracted in CLI.[/yellow]")
        console.print("[yellow]Use an archive manager to extract contents.[/yellow]")
        input("\nPress Enter to continue...")
        
    except Exception as e:
        console.print(f"[red]Error reading archive file: {e}[/red]")
        input("Press Enter to continue...")

def validate_word_document(path):
    """Validate if a Word document can be read"""
    try:
        # Try to open the document
        doc = docx.Document(path)
        
        # Check if document has any content
        has_text = any(para.text.strip() for para in doc.paragraphs)
        has_tables = len(doc.tables) > 0
        
        return True, "Valid document", has_text, has_tables
        
    except Exception as e:
        error_msg = str(e)
        if "Package not found" in error_msg:
            return False, "Document structure corrupted or encrypted", False, False
        elif "lxml" in error_msg.lower():
            return False, "Missing lxml dependency", False, False
        else:
            return False, f"Document error: {error_msg}", False, False

def interactive_pdf_viewer(path, filename):
    """Interactive PDF viewer with navigation and search"""
    try:
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            total_pages = len(reader.pages)
            
            if total_pages == 0:
                console.print("[red]PDF has no pages[/red]")
                input("Press Enter to continue...")
                return
            
            current_page = 0
            search_term = ""
            search_results = []
            search_index = 0
            
            while True:
                clear_screen()
                console.print(Panel(f"[bold cyan]📄 Interactive PDF Viewer[/bold cyan]", expand=False))
                console.print(f"[yellow]File:[/yellow] {filename}")
                console.print(f"[yellow]Page:[/yellow] {current_page + 1} of {total_pages}")
                
                if search_term:
                    console.print(f"[yellow]Search:[/yellow] '{search_term}' ({len(search_results)} results)")
                
                console.print()
                
                # Display current page
                try:
                    page = reader.pages[current_page]
                    text = page.extract_text()
                    
                    if text.strip():
                        # Highlight search results if searching
                        if search_term and search_results:
                            lines = text.split('\n')
                            for i, line in enumerate(lines):
                                if search_term.lower() in line.lower():
                                    # Highlight the line
                                    highlighted = line.replace(
                                        search_term, f"[bold red]{search_term}[/bold red]"
                                    )
                                    lines[i] = highlighted
                            text = '\n'.join(lines)
                        
                        # Display text with pagination
                        lines = text.split('\n')
                        start_line = 0
                        lines_per_page = 20
                        
                        while start_line < len(lines):
                            console.print(f"[bold]Page {current_page + 1} (lines {start_line + 1}-{min(start_line + lines_per_page, len(lines))}):[/bold]")
                            console.print()
                            
                            for i in range(start_line, min(start_line + lines_per_page, len(lines))):
                                console.print(lines[i])
                            
                            if start_line + lines_per_page < len(lines):
                                console.print()
                                console.print("[dim]Press Enter for more lines, 'q' to quit, or use navigation keys[/dim]")
                                key = get_key()
                                if key == 'q':
                                    return
                                elif key == '\r':  # Enter
                                    start_line += lines_per_page
                                    continue
                                else:
                                    break
                            else:
                                break
                    else:
                        console.print("[yellow]No text content found on this page[/yellow]")
                        console.print("[dim]This page might contain only images or be blank[/dim]")
                
                except Exception as e:
                    console.print(f"[red]Error reading page: {e}[/red]")
                
                console.print()
                console.print("[cyan]Navigation:[/cyan]")
                console.print("  [yellow]←/→[/yellow] Previous/Next page")
                console.print("  [yellow]Home/End[/yellow] First/Last page")
                console.print("  [yellow]/[/yellow] Search in document")
                console.print("  [yellow]n[/yellow] Next search result")
                console.print("  [yellow]p[/yellow] Previous search result")
                console.print("  [yellow]q[/yellow] Quit viewer")
                
                key = get_key()
                
                if key == 'q':
                    break
                elif key == '\x1b[D' or key == 'h':  # Left arrow or h
                    current_page = max(0, current_page - 1)
                elif key == '\x1b[C' or key == 'l':  # Right arrow or l
                    current_page = min(total_pages - 1, current_page + 1)
                elif key == '\x1b[H' or key == 'g':  # Home or g
                    current_page = 0
                elif key == '\x1b[F' or key == 'G':  # End or G
                    current_page = total_pages - 1
                elif key == '/':
                    search_term = Prompt.ask("[yellow]Enter search term[/yellow]").strip()
                    if search_term:
                        search_results = []
                        for i, page in enumerate(reader.pages):
                            try:
                                text = page.extract_text()
                                if search_term.lower() in text.lower():
                                    search_results.append(i)
                            except:
                                pass
                        search_index = 0
                        if search_results:
                            current_page = search_results[0]
                        else:
                            console.print(f"[yellow]No results found for '{search_term}'[/yellow]")
                            input("Press Enter to continue...")
                elif key == 'n' and search_results:
                    search_index = (search_index + 1) % len(search_results)
                    current_page = search_results[search_index]
                elif key == 'p' and search_results:
                    search_index = (search_index - 1) % len(search_results)
                    current_page = search_results[search_index]
    
    except Exception as e:
        console.print(f"[red]Error reading PDF: {e}[/red]")
        console.print()
        console.print("[yellow]This might be due to:[/yellow]")
        console.print("• Corrupted or encrypted PDF")
        console.print("• Unsupported PDF version")
        console.print("• PDF contains only images")
        console.print()
        console.print("[cyan]Try updating PyPDF2: pip install --upgrade PyPDF2[/cyan]")
        input("Press Enter to continue...")

def interactive_docx_viewer(path, filename):
    """Interactive DOCX viewer with navigation and search"""
    try:
        doc = docx.Document(path)
        paragraphs = [para for para in doc.paragraphs if para.text.strip()]
        tables = doc.tables
        
        if not paragraphs and not tables:
            console.print("[yellow]Document appears to be empty[/yellow]")
            input("Press Enter to continue...")
            return
        
        current_para = 0
        current_table = 0
        view_mode = "paragraphs"  # "paragraphs" or "tables"
        search_term = ""
        search_results = []
        search_index = 0
        
        while True:
            clear_screen()
            console.print(Panel(f"[bold cyan]📝 Interactive DOCX Viewer[/bold cyan]", expand=False))
            console.print(f"[yellow]File:[/yellow] {filename}")
            
            if view_mode == "paragraphs":
                console.print(f"[yellow]Paragraph:[/yellow] {current_para + 1} of {len(paragraphs)}")
            else:
                console.print(f"[yellow]Table:[/yellow] {current_table + 1} of {len(tables)}")
            
            if search_term:
                console.print(f"[yellow]Search:[/yellow] '{search_term}' ({len(search_results)} results)")
            
            console.print()
            
            # Display current content
            if view_mode == "paragraphs" and paragraphs:
                para = paragraphs[current_para]
                text = para.text
                
                # Highlight search results
                if search_term and search_term.lower() in text.lower():
                    text = text.replace(
                        search_term, f"[bold red]{search_term}[/bold red]"
                    )
                
                console.print(f"[bold]Paragraph {current_para + 1}:[/bold]")
                console.print(text)
                
            elif view_mode == "tables" and tables:
                table = tables[current_table]
                console.print(f"[bold]Table {current_table + 1}:[/bold]")
                console.print()
                
                # Display table in a simple format
                for i, row in enumerate(table.rows):
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()[:20]  # Limit cell width
                        if len(cell.text.strip()) > 20:
                            cell_text += "..."
                        row_text.append(cell_text)
                    console.print(f"Row {i+1}: {' | '.join(row_text)}")
            
            console.print()
            console.print("[cyan]Navigation:[/cyan]")
            console.print("  [yellow]←/→[/yellow] Previous/Next paragraph/table")
            console.print("  [yellow]t[/yellow] Toggle between paragraphs and tables")
            console.print("  [yellow]Home/End[/yellow] First/Last paragraph/table")
            console.print("  [yellow]/[/yellow] Search in document")
            console.print("  [yellow]n[/yellow] Next search result")
            console.print("  [yellow]p[/yellow] Previous search result")
            console.print("  [yellow]q[/yellow] Quit viewer")
            
            key = get_key()
            
            if key == 'q':
                break
            elif key == '\x1b[D' or key == 'h':  # Left arrow or h
                if view_mode == "paragraphs":
                    current_para = max(0, current_para - 1)
                else:
                    current_table = max(0, current_table - 1)
            elif key == '\x1b[C' or key == 'l':  # Right arrow or l
                if view_mode == "paragraphs":
                    current_para = min(len(paragraphs) - 1, current_para + 1)
                else:
                    current_table = min(len(tables) - 1, current_table + 1)
            elif key == '\x1b[H' or key == 'g':  # Home or g
                if view_mode == "paragraphs":
                    current_para = 0
                else:
                    current_table = 0
            elif key == '\x1b[F' or key == 'G':  # End or G
                if view_mode == "paragraphs":
                    current_para = len(paragraphs) - 1
                else:
                    current_table = len(tables) - 1
            elif key == 't':
                view_mode = "tables" if view_mode == "paragraphs" else "paragraphs"
                if view_mode == "paragraphs":
                    current_para = 0
                else:
                    current_table = 0
            elif key == '/':
                search_term = Prompt.ask("[yellow]Enter search term[/yellow]").strip()
                if search_term:
                    search_results = []
                    for i, para in enumerate(paragraphs):
                        if search_term.lower() in para.text.lower():
                            search_results.append(("paragraph", i))
                    for i, table in enumerate(tables):
                        for row in table.rows:
                            for cell in row.cells:
                                if search_term.lower() in cell.text.lower():
                                    search_results.append(("table", i))
                                    break
                    
                    search_index = 0
                    if search_results:
                        result_type, result_index = search_results[0]
                        if result_type == "paragraph":
                            view_mode = "paragraphs"
                            current_para = result_index
                        else:
                            view_mode = "tables"
                            current_table = result_index
                    else:
                        console.print(f"[yellow]No results found for '{search_term}'[/yellow]")
                        input("Press Enter to continue...")
            elif key == 'n' and search_results:
                search_index = (search_index + 1) % len(search_results)
                result_type, result_index = search_results[search_index]
                if result_type == "paragraph":
                    view_mode = "paragraphs"
                    current_para = result_index
                else:
                    view_mode = "tables"
                    current_table = result_index
            elif key == 'p' and search_results:
                search_index = (search_index - 1) % len(search_results)
                result_type, result_index = search_results[search_index]
                if result_type == "paragraph":
                    view_mode = "paragraphs"
                    current_para = result_index
                else:
                    view_mode = "tables"
                    current_table = result_index
    
    except Exception as e:
        console.print(f"[red]Error reading DOCX: {e}[/red]")
        input("Press Enter to continue...")

def interactive_csv_viewer(path, filename):
    """Interactive CSV viewer with navigation and search"""
    try:
        with open(path, "r", newline="", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            rows = list(reader)
            
            if not rows:
                console.print("[yellow]CSV file is empty[/yellow]")
                input("Press Enter to continue...")
                return
            
            current_row = 0
            rows_per_page = 15
            search_term = ""
            search_results = []
            search_index = 0
            
            while True:
                clear_screen()
                console.print(Panel(f"[bold cyan]📊 Interactive CSV Viewer[/bold cyan]", expand=False))
                console.print(f"[yellow]File:[/yellow] {filename}")
                console.print(f"[yellow]Total rows:[/yellow] {len(rows)}")
                console.print(f"[yellow]Showing rows:[/yellow] {current_row + 1}-{min(current_row + rows_per_page, len(rows))}")
                
                if search_term:
                    console.print(f"[yellow]Search:[/yellow] '{search_term}' ({len(search_results)} results)")
                
                console.print()
                
                # Display current rows
                for i in range(current_row, min(current_row + rows_per_page, len(rows))):
                    row = rows[i]
                    row_text = []
                    
                    for cell in row:
                        # Limit cell width and highlight search results
                        cell_text = str(cell)[:15]
                        if len(str(cell)) > 15:
                            cell_text += "..."
                        
                        if search_term and search_term.lower() in str(cell).lower():
                            cell_text = cell_text.replace(
                                search_term, f"[bold red]{search_term}[/bold red]"
                            )
                        
                        row_text.append(cell_text)
                    
                    # Highlight row number if it's a search result
                    row_num = f"Row {i+1}"
                    if search_term and any(search_term.lower() in str(cell).lower() for cell in row):
                        row_num = f"[bold red]Row {i+1}[/bold red]"
                    
                    console.print(f"{row_num}: {' | '.join(row_text)}")
                
                console.print()
                console.print("[cyan]Navigation:[/cyan]")
                console.print("  [yellow]↑/↓[/yellow] Previous/Next page")
                console.print("  [yellow]Home/End[/yellow] First/Last page")
                console.print("  [yellow]/[/yellow] Search in CSV")
                console.print("  [yellow]n[/yellow] Next search result")
                console.print("  [yellow]p[/yellow] Previous search result")
                console.print("  [yellow]q[/yellow] Quit viewer")
                
                key = get_key()
                
                if key == 'q':
                    break
                elif key == '\x1b[A' or key == 'k':  # Up arrow or k
                    current_row = max(0, current_row - rows_per_page)
                elif key == '\x1b[B' or key == 'j':  # Down arrow or j
                    current_row = min(len(rows) - rows_per_page, current_row + rows_per_page)
                elif key == '\x1b[H' or key == 'g':  # Home or g
                    current_row = 0
                elif key == '\x1b[F' or key == 'G':  # End or G
                    current_row = max(0, len(rows) - rows_per_page)
                elif key == '/':
                    search_term = Prompt.ask("[yellow]Enter search term[/yellow]").strip()
                    if search_term:
                        search_results = []
                        for i, row in enumerate(rows):
                            for cell in row:
                                if search_term.lower() in str(cell).lower():
                                    search_results.append(i)
                                    break
                        search_index = 0
                        if search_results:
                            current_row = (search_results[0] // rows_per_page) * rows_per_page
                        else:
                            console.print(f"[yellow]No results found for '{search_term}'[/yellow]")
                            input("Press Enter to continue...")
                elif key == 'n' and search_results:
                    search_index = (search_index + 1) % len(search_results)
                    current_row = (search_results[search_index] // rows_per_page) * rows_per_page
                elif key == 'p' and search_results:
                    search_index = (search_index - 1) % len(search_results)
                    current_row = (search_results[search_index] // rows_per_page) * rows_per_page
    
    except Exception as e:
        console.print(f"[red]Error reading CSV: {e}[/red]")
        input("Press Enter to continue...")

def view_file_rich(user_folder, subject, filename, editable=False):
    """Main file viewer function - handles ALL file types in CLI"""
    path = os.path.join("subjects", user_folder, subject, filename) if user_folder else os.path.join("subjects", subject, filename)
    
    if not os.path.exists(path):
        console.print(f"[red]File '{filename}' not found[/red]")
        return
    
    try:
        ext = filename.split(".")[-1].lower() if "." in filename else "txt"
        console.print(Panel(f"[bold cyan]Opening {filename} in CLI[/bold cyan]", expand=False))
        
        if editable:
            edit_file_with_reason(user_folder, subject, filename)
            return

        # Text-based files (interactive reader)
        if ext in ["txt", "md", "py", "json", "js", "html", "css", "xml", "yaml", "yml", "ini", "cfg", "conf", "log", "sh", "bat", "ps1"]:
            interactive_text_reader(path, filename)
            
        # PDF files
        elif ext == "pdf":
            if not PyPDF2:
                console.print("[red]PyPDF2 not installed - showing file info instead[/red]")
                file_size = os.path.getsize(path)
                console.print(f"[cyan]PDF file: {filename}[/cyan]")
                console.print(f"[cyan]Size: {file_size:,} bytes[/cyan]")
                console.print()
                console.print("[yellow]To view PDF content:[/yellow]")
                console.print("[cyan]pip install PyPDF2[/cyan]")
                input("\nPress Enter to continue...")
                return
            
            # Interactive PDF viewer
            interactive_pdf_viewer(path, filename)

        # Word documents
        elif ext in ["doc", "docx"]:
            if not docx:
                console.print("[red]python-docx not installed - showing file info instead[/red]")
                file_size = os.path.getsize(path)
                console.print(f"[cyan]Word document: {filename}[/cyan]")
                console.print(f"[cyan]Size: {file_size:,} bytes[/cyan]")
                console.print()
                console.print("[yellow]To view Word document content:[/yellow]")
                console.print("[cyan]pip install python-docx lxml[/cyan]")
                console.print()
                console.print("[dim]Note: lxml is required for python-docx to work properly[/dim]")
                input("\nPress Enter to continue...")
                return
            
            # Validate document first
            console.print("[yellow]🔍 Validating document...[/yellow]")
            is_valid, message, has_text, has_tables = validate_word_document(path)
            
            if not is_valid:
                console.print(f"[red]❌ Document validation failed: {message}[/red]")
                console.print()
                
                # Provide specific error diagnosis
                if "corrupted or encrypted" in message:
                    console.print("[yellow]🔍 Diagnosis: Document structure issue[/yellow]")
                    console.print("This usually means:")
                    console.print("• Document is corrupted or damaged")
                    console.print("• Document is encrypted or password-protected")
                    console.print("• Document contains unsupported elements")
                    console.print("• Document was created with a very old/new version of Word")
                elif "lxml" in message.lower():
                    console.print("[yellow]🔍 Diagnosis: Missing XML processing library[/yellow]")
                    console.print("• lxml is required for python-docx to work properly")
                else:
                    console.print("[yellow]🔍 Diagnosis: General document reading error[/yellow]")
                    console.print("This might be due to:")
                    console.print("• Corrupted or encrypted document")
                    console.print("• Unsupported Word format")
                    console.print("• Missing dependencies")
                
                console.print()
                console.print("[cyan]💡 Solutions to try:[/cyan]")
                console.print("1. [yellow]Install/update lxml:[/yellow] pip install --upgrade lxml")
                console.print("2. [yellow]Try opening in Word and save as .docx again[/yellow]")
                console.print("3. [yellow]Convert to PDF and upload the PDF instead[/yellow]")
                console.print("4. [yellow]Check if document is password-protected[/yellow]")
                console.print("5. [yellow]Try uploading a different Word document[/yellow]")
                
                # Show file info as fallback
                try:
                    file_size = os.path.getsize(path)
                    console.print()
                    console.print(f"[cyan]File info:[/cyan] {filename}")
                    console.print(f"[cyan]Size:[/cyan] {file_size:,} bytes")
                    console.print(f"[cyan]Path:[/cyan] {path}")
                except:
                    pass
                
                input("\nPress Enter to continue...")
                return
            
            # Document is valid, use interactive viewer
            console.print(f"[green]✅ Document validated successfully[/green]")
            console.print(f"[cyan]Document has {len(doc.paragraphs)} paragraphs[/cyan]")
            
            if has_tables:
                console.print(f"[cyan]Document has {len(doc.tables)} tables[/cyan]")
            
            console.print()
            console.print("[yellow]Opening interactive DOCX viewer...[/yellow]")
            input("Press Enter to continue...")
            
            # Use interactive DOCX viewer
            interactive_docx_viewer(path, filename)

        # CSV files
        elif ext == "csv":
            interactive_csv_viewer(path, filename)

        # Image files
        elif ext in ["jpg", "jpeg", "png", "gif", "bmp", "tiff", "svg", "webp"]:
            view_image_info(path, filename, ext)

        # Archive files
        elif ext in ["zip", "rar", "7z", "tar", "gz", "bz2", "xz"]:
            view_archive_info(path, filename, ext)

        # Executable files
        elif ext in ["exe", "dll", "so", "dylib", "bin"]:
            console.print(f"[yellow]Executable file: {filename}[/yellow]")
            console.print("[yellow]Cannot execute or view binary executables in CLI.[/yellow]")
            input("Press Enter to continue...")

        # Audio/Video files
        elif ext in ["mp3", "wav", "flac", "mp4", "avi", "mkv", "mov", "wmv"]:
            file_size = os.path.getsize(path)
            console.print(f"[cyan]Media file: {filename}[/cyan]")
            console.print(f"[cyan]Format: {ext.upper()}[/cyan]")
            console.print(f"[cyan]Size: {file_size:,} bytes[/cyan]")
            console.print("[yellow]Media files cannot be played in CLI.[/yellow]")
            input("\nPress Enter to continue...")

        # Default: try to read as text, fallback to binary
        else:
            try:
                # First try to read as text
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read(1000)  # Read first 1000 chars
                    if content.strip():
                        console.print(f"[yellow]Unknown text format '{ext}' - showing as text:[/yellow]")
                        console.print(content)
                        if len(content) == 1000:
                            console.print("\n[yellow]... (showing first 1000 characters)[/yellow]")
                        input("\nPress Enter to continue...")
                    else:
                        raise UnicodeDecodeError("", b"", 0, 1, "empty content")
            except (UnicodeDecodeError, UnicodeError):
                # If text reading fails, show as binary
                console.print(f"[yellow]Unknown binary format '{ext}' - showing hex dump:[/yellow]")
                view_binary_file(path, filename, ext)

    except Exception as e:
        console.print(f"[red]Unexpected error opening file: {e}[/red]")
        console.print("[yellow]File may be corrupted or in an unsupported format[/yellow]")
        input("Press Enter to continue...")

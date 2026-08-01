# file_viewer.py
import os, csv, sys, subprocess, tempfile
from rich.console import Console
from rich.table import Table
from rich.text import Text
from rich.prompt import Prompt
from datetime import datetime
from study_cli_hub.animations import cli_panel as Panel
from study_cli_hub.error_handler import handle_error
from study_cli_hub.paths import note_path

console = Console()

TEXT_EXTENSIONS = {
    "txt", "md", "py", "json", "js", "html", "css", "xml", "yaml", "yml",
    "ini", "cfg", "conf", "log", "sh", "bat", "ps1",
}

try:
    import PyPDF2
except:
    PyPDF2 = None

try:
    import docx
except:
    docx = None

_ZIP_SIGNATURES = (b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")


def detect_binary_content(path, sample_size=8192):
    """Best-effort check for whether a file's content is actually text, so a
    misnamed/mislabeled file (e.g. a .numbers or .xlsx saved with a .csv
    extension) shows a clear message instead of being rendered as garbled
    text. Returns None if the content looks like real text, otherwise a
    short human-readable reason string."""
    try:
        with open(path, "rb") as f:
            sample = f.read(sample_size)
    except OSError:
        return None

    if not sample:
        return None

    if sample.startswith(_ZIP_SIGNATURES):
        return "it looks like a ZIP-based file (e.g. .xlsx, .numbers, .docx) saved with the wrong extension"
    if sample.startswith(b"%PDF-"):
        return "it looks like a PDF file saved with the wrong extension"
    if b"\x00" in sample:
        return "it contains null bytes, which real text files don't have"

    # A real text file should decode cleanly as UTF-8 (or very close to it).
    try:
        sample.decode("utf-8")
        return None
    except UnicodeDecodeError:
        pass

    non_printable = sum(1 for b in sample if b < 9 or (13 < b < 32))
    if non_printable / len(sample) > 0.05:
        return "it contains a high proportion of non-text bytes"
    return None


def clear_screen():
    """Clear the terminal screen"""
    os.system("cls" if os.name == "nt" else "clear")


def open_in_editor(initial_content=""):
    """Open $VISUAL/$EDITOR on a temp file pre-filled with initial_content
    and return the edited text - the same pattern git/crontab use for
    editing. Falls back to a plain type-then-Ctrl+D prompt when no editor
    is configured."""
    editor = os.environ.get("VISUAL") or os.environ.get("EDITOR")
    if editor:
        fd, tmp_path = tempfile.mkstemp(suffix=".md")
        try:
            with os.fdopen(fd, "w", encoding="utf-8") as f:
                f.write(initial_content)
            console.print(f"[dim]Opening in {editor}...[/dim]")
            subprocess.run([editor, tmp_path])
            with open(tmp_path, "r", encoding="utf-8") as f:
                return f.read()
        finally:
            try:
                os.remove(tmp_path)
            except OSError:
                pass

    console.print("[dim]No $EDITOR/$VISUAL set - falling back to inline entry "
                   "(set $EDITOR to use vim/nano/etc. instead).[/dim]")
    console.print("[yellow]Type your content, then press Ctrl+D when done:[/yellow]")
    if initial_content:
        console.print(initial_content)
        console.print("[dim]--- existing content shown above; typed lines below replace it ---[/dim]")
    lines = []
    try:
        while True:
            lines.append(input())
    except EOFError:
        pass
    return "\n".join(lines) if lines else initial_content


def print_nav_table(items, extra=None):
    """Renders a viewer's key bindings as a compact table, meant to be
    printed last so it sits pinned at the bottom of the screen right above
    the key-press prompt. `items` is a list of (key, action) pairs; `extra`
    is an optional trailing status line (e.g. active search term)."""
    table = Table(show_header=False, box=None, padding=(0, 1, 0, 0))
    table.add_column(style="bold yellow", no_wrap=True)
    table.add_column(style="dim")
    for key, action in items:
        table.add_row(key, action)
    console.print(table)
    if extra:
        console.print(extra)

def get_key():
    """Cross-platform key input with improved navigation"""
    try:
        import msvcrt
        if msvcrt.kbhit():
            key = msvcrt.getch()
            if key == b'\xe0':  # Special key
                key = msvcrt.getch()
                if key == b'H': return '\x1b[A'  # Up arrow
                elif key == b'P': return '\x1b[B'  # Down arrow
                elif key == b'K': return '\x1b[D'  # Left arrow
                elif key == b'M': return '\x1b[C'  # Right arrow
                elif key == b'G': return '\x1b[H'  # Home
                elif key == b'O': return '\x1b[F'  # End
                elif key == b'I': return '\x1b[5~'  # Page Up
                elif key == b'Q': return '\x1b[6~'  # Page Down
            elif key == b' ': return ' '
            elif key == b'q': return 'q'
            elif key == b'/': return '/'
            elif key == b':': return ':'
            elif key == b'h': return 'h'
            elif key == b'n': return 'n'
            elif key == b'p': return 'p'
            elif key == b't': return 't'
            elif key == b'g': return 'g'
            elif key == b'G': return 'G'
            elif key == b'j': return 'j'
            elif key == b'k': return 'k'
            elif key == b'l': return 'l'
            elif key == b'?': return '?'
            elif key == b'\r': return '\r'
    except ImportError:
        # Unix/Linux/Mac
        import termios, tty
        if not sys.stdin.isatty():
            return None
        fd = sys.stdin.fileno()
        old_settings = termios.tcgetattr(fd)
        try:
            tty.setraw(fd)
            ch = sys.stdin.read(1)
            if ch == '\x1b':  # ESC sequence
                seq = sys.stdin.read(1)
                if seq != '[':
                    return None
                third = sys.stdin.read(1)
                if third in 'ABCDHF':
                    # 3-byte sequences: ESC [ A/B/C/D/H/F
                    return '\x1b[' + third
                elif third in '56':
                    # 4-byte sequences: ESC [ 5/6 ~  (Page Up/Down)
                    fourth = sys.stdin.read(1)
                    if fourth == '~':
                        return '\x1b[' + third + fourth
                return None
            elif ch == ' ': return ' '
            elif ch == 'q': return 'q'
            elif ch == '/': return '/'
            elif ch == ':': return ':'
            elif ch == 'h': return 'h'
            elif ch == 'n': return 'n'
            elif ch == 'p': return 'p'
            elif ch == 't': return 't'
            elif ch == 'g': return 'g'
            elif ch == 'G': return 'G'
            elif ch == 'j': return 'j'
            elif ch == 'k': return 'k'
            elif ch == 'l': return 'l'
            elif ch == '?': return '?'
            elif ch == '\r': return '\r'
        finally:
            termios.tcsetattr(fd, termios.TCSADRAIN, old_settings)
    return None

def interactive_text_reader(path, filename, start_line=0):
    """Interactive text file reader with navigation and highlighting"""
    reason = detect_binary_content(path)
    if reason:
        console.print(Panel(
            f"[bold red]'{filename}' doesn't look like a real text file[/bold red]\n\n"
            f"{reason}.\n\n"
            "[dim]Use /repair or re-upload it with the correct extension to "
            "view it properly.[/dim]",
            expand=False,
        ))
        input("Press Enter to continue...")
        return

    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            lines = f.readlines()

        highlighted = set()
        search_results = []
        search_index = 0
        current_line = max(0, min(start_line, len(lines) - 1)) if lines else 0
        lines_per_page = 20

        while True:
            console.clear()

            # Header
            header = f"[bold cyan]{filename}[/bold cyan] | Line {current_line + 1}/{len(lines)}"
            if search_results:
                header += f" | Search: {len(search_results)} results"
            console.print(Panel(header, expand=False))
            console.print()

            # Display lines
            start_line = max(0, current_line - lines_per_page // 2)
            end_line = min(len(lines), start_line + lines_per_page)

            for i in range(start_line, end_line):
                line_no = f"{i+1:>4}"
                line_text = lines[i].rstrip()

                # Highlighting
                if i in highlighted:
                    console.print(f"[bold black on yellow]{line_no} {line_text}[/bold black on yellow]")
                elif i == current_line:
                    console.print(f"[bold white on blue]{line_no} {line_text}[/bold white on blue]")
                else:
                    console.print(f"[green]{line_no}[/green] {line_text}")

            # Footer
            if current_line >= len(lines) - 1:
                console.print("\n[bold yellow]End of file[/bold yellow]")

            console.print()
            print_nav_table([
                ("↑/↓ k/j", "Scroll line by line"),
                ("PgUp/PgDn u/d", "Jump a page"),
                ("g / G", "Go to beginning / end"),
                (":", "Go to a specific line number"),
                ("Space / h", "Highlight/unhighlight line"),
                ("/", "Search in file"),
                ("n / p", "Next / previous search result"),
                ("?", "Help"),
                ("q", "Quit"),
            ])

            # Get user input
            key = get_key()
            if not key:
                continue
                
            if key == '\x1b[A' or key == 'k':  # Up arrow or k
                current_line = max(0, current_line - 1)
            elif key == '\x1b[B' or key == 'j':  # Down arrow or j
                current_line = min(len(lines) - 1, current_line + 1)
            elif key == '\x1b[5~' or key == 'u':  # Page Up or u
                current_line = max(0, current_line - lines_per_page)
            elif key == '\x1b[6~' or key == 'd':  # Page Down or d
                current_line = min(len(lines) - 1, current_line + lines_per_page)
            elif key == ' ' or key == 'h':  # Space or h for highlight
                # Toggle highlight for current line
                if current_line in highlighted:
                    highlighted.remove(current_line)
                else:
                    highlighted.add(current_line)
            elif key == '/':  # Search
                search_term = Prompt.ask("[yellow]Enter search term[/yellow]").strip()
                if search_term:
                    search_results = []
                    for i, line in enumerate(lines):
                        if search_term.lower() in line.lower():
                            search_results.append(i)
                    if search_results:
                        current_line = search_results[0]
                        console.print(f"[green]Found {len(search_results)} results[/green]")
                    else:
                        console.print("[red]No results found[/red]")
                        input("Press Enter to continue...")
            elif key == ':':  # Go to line
                target = Prompt.ask("[yellow]Go to line number[/yellow]").strip()
                if target.isdigit() and 1 <= int(target) <= len(lines):
                    current_line = int(target) - 1
                else:
                    console.print(f"[red]Invalid line number! Choose 1-{len(lines)}[/red]")
                    input("Press Enter to continue...")
            elif key == '?':  # Help
                show_reader_help()
            elif key == 'q':  # Quit
                break
            elif key == 'n' and search_results:  # Next search result
                search_index = (search_index + 1) % len(search_results)
                current_line = search_results[search_index]
            elif key == 'p' and search_results:  # Previous search result
                search_index = (search_index - 1) % len(search_results)
                current_line = search_results[search_index]
            elif key == 'g':  # Go to beginning
                current_line = 0
            elif key == 'G':  # Go to end
                current_line = len(lines) - 1
                
    except Exception as e:
        handle_error(e)

def show_reader_help():
    """Show help for the interactive reader"""
    console.clear()
    console.print(Panel("[bold cyan]Interactive Reader Help[/bold cyan]", expand=False))
    console.print()
    print_nav_table([
        ("↑/↓ k/j", "Scroll line by line"),
        ("PgUp/PgDn u/d", "Jump a page"),
        ("g / G", "Go to beginning / end"),
        (":", "Go to a specific line number"),
        ("Space / h", "Highlight/unhighlight current line"),
        ("/", "Search for text in file"),
        ("n / p", "Next / previous search result"),
        ("?", "Show this help"),
        ("q", "Quit reader"),
    ])
    console.print()
    input("Press Enter to continue...")

def edit_file_with_reason(user_folder, subject, filename):
    """Edit a file with reason tracking"""
    path = note_path(user_folder, subject, filename)

    if not os.path.exists(path):
        console.print(f"[red]File '{filename}' not found[/red]")
        return

    try:
        # Get edit reason
        reason = Prompt.ask("[yellow]Enter reason for editing this file[/yellow]").strip()
        if not reason:
            reason = "No reason provided"
        
        # Read current content
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        
        # Show current content
        console.print(Panel(f"[bold cyan]Editing: {filename}[/bold cyan]", expand=False))
        console.print(f"[dim]Reason: {reason}[/dim]")
        console.print()

        # Edit (via $EDITOR if set, otherwise inline) - pre-filled with the
        # existing content so you're editing in place, not retyping the file
        new_content = open_in_editor(content)

        # Save with reason tracking
        backup_path = path + f".backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        with open(backup_path, "w", encoding="utf-8") as f:
            f.write(content)
        
        with open(path, "w", encoding="utf-8") as f:
            f.write(new_content)
        
        # Log the edit
        log_path = note_path(user_folder, subject, "edit_log.txt")
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(f"[{datetime.now()}] Edited {filename} - Reason: {reason}\n")
            f.write(f"Backup saved as: {os.path.basename(backup_path)}\n\n")
        
        console.print(f"[green]✅ File edited successfully![/green]")
        console.print(f"[dim]Backup saved as: {os.path.basename(backup_path)}[/dim]")
        
    except Exception as e:
        handle_error(e)

def view_binary_file(path, filename, ext):
    """View binary files as hex dump"""
    try:
        console.print(f"[yellow]Binary file detected. Showing hex dump:[/yellow]")
        console.print()
        
        with open(path, "rb") as f:
            chunk_size = 16
            offset = 0
            
            for chunk in iter(lambda: f.read(chunk_size), b""):
                # Create hex representation
                hex_str = " ".join(f"{b:02x}" for b in chunk)
                hex_str = hex_str.ljust(47)  # Pad to align ASCII
                
                # Create ASCII representation
                ascii_str = "".join(chr(b) if 32 <= b <= 126 else "." for b in chunk)
                
                console.print(f"{offset:08x}: {hex_str} |{ascii_str}|")
                offset += len(chunk)
                
                # Limit output for very large files
                if offset > 1024:  # Show first 1KB
                    console.print(f"\n[yellow]... (showing first 1KB of {os.path.getsize(path)} bytes)[/yellow]")
                    break
        
        input("\nPress Enter to continue...")
        
    except Exception as e:
        console.print(f"[red]Error reading binary file: {e}[/red]")
        input("Press Enter to continue...")

def view_image_info(path, filename, ext):
    """Show image file information"""
    try:
        file_size = os.path.getsize(path)
        console.print(f"[cyan]Image file: {filename}[/cyan]")
        console.print(f"[cyan]Format: {ext.upper()}[/cyan]")
        console.print(f"[cyan]Size: {file_size:,} bytes[/cyan]")
        console.print()
        console.print("[yellow]Image files cannot be displayed in CLI.[/yellow]")
        console.print("[yellow]Use an image viewer or convert to text format.[/yellow]")
        input("\nPress Enter to continue...")
        
    except Exception as e:
        console.print(f"[red]Error reading image file: {e}[/red]")
        input("Press Enter to continue...")

def view_archive_info(path, filename, ext):
    """Show archive file information"""
    try:
        file_size = os.path.getsize(path)
        console.print(f"[cyan]Archive file: {filename}[/cyan]")
        console.print(f"[cyan]Format: {ext.upper()}[/cyan]")
        console.print(f"[cyan]Size: {file_size:,} bytes[/cyan]")
        console.print()
        console.print("[yellow]Archive files cannot be extracted in CLI.[/yellow]")
        console.print("[yellow]Use an archive manager to extract contents.[/yellow]")
        input("\nPress Enter to continue...")
        
    except Exception as e:
        console.print(f"[red]Error reading archive file: {e}[/red]")
        input("Press Enter to continue...")

def validate_word_document(path):
    """Validate if a Word document can be read"""
    try:
        # Try to open the document
        doc = docx.Document(path)
        
        # Check if document has any content
        has_text = any(para.text.strip() for para in doc.paragraphs)
        has_tables = len(doc.tables) > 0
        
        return True, "Valid document", has_text, has_tables
        
    except Exception as e:
        error_msg = str(e)
        if "Package not found" in error_msg:
            return False, "Document structure corrupted or encrypted", False, False
        elif "lxml" in error_msg.lower():
            return False, "Missing lxml dependency", False, False
        else:
            return False, f"Document error: {error_msg}", False, False

def interactive_pdf_viewer(path, filename, start_page=0):
    """Interactive PDF viewer with navigation and search"""
    try:
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            total_pages = len(reader.pages)

            if total_pages == 0:
                console.print("[red]PDF has no pages[/red]")
                input("Press Enter to continue...")
                return

            current_page = max(0, min(start_page, total_pages - 1))
            search_term = ""
            search_results = []
            search_index = 0
            
            while True:
                clear_screen()
                console.print(Panel(f"[bold cyan]📄 Interactive PDF Viewer[/bold cyan]", expand=False))
                console.print(f"[yellow]File:[/yellow] {filename}")
                console.print(f"[yellow]Page:[/yellow] {current_page + 1} of {total_pages}")
                
                if search_term:
                    console.print(f"[yellow]Search:[/yellow] '{search_term}' ({len(search_results)} results)")
                
                console.print()
                
                # Display current page
                try:
                    page = reader.pages[current_page]
                    text = page.extract_text()
                    
                    if text.strip():
                        # Highlight search results if searching
                        if search_term and search_results:
                            lines = text.split('\n')
                            for i, line in enumerate(lines):
                                if search_term.lower() in line.lower():
                                    # Highlight the line
                                    highlighted = line.replace(
                                        search_term, f"[bold red]{search_term}[/bold red]"
                                    )
                                    lines[i] = highlighted
                            text = '\n'.join(lines)
                        
                        # Display text with pagination
                        lines = text.split('\n')
                        start_line = 0
                        lines_per_page = 20
                        
                        while start_line < len(lines):
                            console.print(f"[bold]Page {current_page + 1} (lines {start_line + 1}-{min(start_line + lines_per_page, len(lines))}):[/bold]")
                            console.print()
                            
                            for i in range(start_line, min(start_line + lines_per_page, len(lines))):
                                console.print(lines[i])
                            
                            if start_line + lines_per_page < len(lines):
                                console.print()
                                console.print("[dim]Press Enter for more lines, 'q' to quit, or use navigation keys[/dim]")
                                key = get_key()
                                if key == 'q':
                                    return
                                elif key == '\r':  # Enter
                                    start_line += lines_per_page
                                    continue
                                else:
                                    break
                            else:
                                break
                    else:
                        console.print("[yellow]No text content found on this page[/yellow]")
                        console.print("[dim]This page might contain only images or be blank[/dim]")
                
                except Exception as e:
                    console.print(f"[red]Error reading page: {e}[/red]")
                
                console.print()
                print_nav_table([
                    ("←/→ h/l", "Previous / next page"),
                    ("Home/End g/G", "First / last page"),
                    (":", "Go to a specific page number"),
                    ("/", "Search in document"),
                    ("n / p", "Next / previous search result"),
                    ("q", "Quit"),
                ])

                key = get_key()

                if key == 'q':
                    break
                elif key == '\x1b[D' or key == 'h':  # Left arrow or h
                    current_page = max(0, current_page - 1)
                elif key == '\x1b[C' or key == 'l':  # Right arrow or l
                    current_page = min(total_pages - 1, current_page + 1)
                elif key == '\x1b[H' or key == 'g':  # Home or g
                    current_page = 0
                elif key == '\x1b[F' or key == 'G':  # End or G
                    current_page = total_pages - 1
                elif key == ':':
                    target = Prompt.ask("[yellow]Go to page number[/yellow]").strip()
                    if target.isdigit() and 1 <= int(target) <= total_pages:
                        current_page = int(target) - 1
                    else:
                        console.print(f"[red]Invalid page number! Choose 1-{total_pages}[/red]")
                        input("Press Enter to continue...")
                elif key == '/':
                    search_term = Prompt.ask("[yellow]Enter search term[/yellow]").strip()
                    if search_term:
                        search_results = []
                        for i, page in enumerate(reader.pages):
                            try:
                                text = page.extract_text()
                                if search_term.lower() in text.lower():
                                    search_results.append(i)
                            except:
                                pass
                        search_index = 0
                        if search_results:
                            current_page = search_results[0]
                        else:
                            console.print(f"[yellow]No results found for '{search_term}'[/yellow]")
                            input("Press Enter to continue...")
                elif key == 'n' and search_results:
                    search_index = (search_index + 1) % len(search_results)
                    current_page = search_results[search_index]
                elif key == 'p' and search_results:
                    search_index = (search_index - 1) % len(search_results)
                    current_page = search_results[search_index]
    
    except Exception as e:
        console.print(f"[red]Error reading PDF: {e}[/red]")
        console.print()
        console.print("[yellow]This might be due to:[/yellow]")
        console.print("• Corrupted or encrypted PDF")
        console.print("• Unsupported PDF version")
        console.print("• PDF contains only images")
        console.print()
        console.print("[cyan]Try updating PyPDF2: pip install --upgrade PyPDF2[/cyan]")
        input("Press Enter to continue...")

def interactive_docx_viewer(path, filename, start_para=0):
    """Interactive DOCX viewer with navigation and search"""
    try:
        doc = docx.Document(path)
        paragraphs = [para for para in doc.paragraphs if para.text.strip()]
        tables = doc.tables

        if not paragraphs and not tables:
            console.print("[yellow]Document appears to be empty[/yellow]")
            input("Press Enter to continue...")
            return

        current_para = max(0, min(start_para, len(paragraphs) - 1)) if paragraphs else 0
        current_table = 0
        view_mode = "paragraphs"  # "paragraphs" or "tables"
        search_term = ""
        search_results = []
        search_index = 0
        
        while True:
            clear_screen()
            console.print(Panel(f"[bold cyan]📝 Interactive DOCX Viewer[/bold cyan]", expand=False))
            console.print(f"[yellow]File:[/yellow] {filename}")
            
            if view_mode == "paragraphs":
                console.print(f"[yellow]Paragraph:[/yellow] {current_para + 1} of {len(paragraphs)}")
            else:
                console.print(f"[yellow]Table:[/yellow] {current_table + 1} of {len(tables)}")
            
            if search_term:
                console.print(f"[yellow]Search:[/yellow] '{search_term}' ({len(search_results)} results)")
            
            console.print()
            
            # Display current content
            if view_mode == "paragraphs" and paragraphs:
                para = paragraphs[current_para]
                text = para.text
                
                # Highlight search results
                if search_term and search_term.lower() in text.lower():
                    text = text.replace(
                        search_term, f"[bold red]{search_term}[/bold red]"
                    )
                
                console.print(f"[bold]Paragraph {current_para + 1}:[/bold]")
                console.print(text)
                
            elif view_mode == "tables" and tables:
                table = tables[current_table]
                console.print(f"[bold]Table {current_table + 1}:[/bold]")
                console.print()
                
                # Display table in a simple format
                for i, row in enumerate(table.rows):
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()[:20]  # Limit cell width
                        if len(cell.text.strip()) > 20:
                            cell_text += "..."
                        row_text.append(cell_text)
                    console.print(f"Row {i+1}: {' | '.join(row_text)}")
            
            console.print()
            print_nav_table([
                ("←/→ h/l", "Previous / next paragraph/table"),
                ("t", "Toggle paragraphs / tables"),
                ("Home/End g/G", "First / last"),
                (":", "Go to a specific paragraph/table number"),
                ("/", "Search in document"),
                ("n / p", "Next / previous search result"),
                ("q", "Quit"),
            ])

            key = get_key()

            if key == 'q':
                break
            elif key == '\x1b[D' or key == 'h':  # Left arrow or h
                if view_mode == "paragraphs":
                    current_para = max(0, current_para - 1)
                else:
                    current_table = max(0, current_table - 1)
            elif key == '\x1b[C' or key == 'l':  # Right arrow or l
                if view_mode == "paragraphs":
                    current_para = min(len(paragraphs) - 1, current_para + 1)
                else:
                    current_table = min(len(tables) - 1, current_table + 1)
            elif key == '\x1b[H' or key == 'g':  # Home or g
                if view_mode == "paragraphs":
                    current_para = 0
                else:
                    current_table = 0
            elif key == '\x1b[F' or key == 'G':  # End or G
                if view_mode == "paragraphs":
                    current_para = len(paragraphs) - 1
                else:
                    current_table = len(tables) - 1
            elif key == ':':
                count = len(paragraphs) if view_mode == "paragraphs" else len(tables)
                label = "paragraph" if view_mode == "paragraphs" else "table"
                target = Prompt.ask(f"[yellow]Go to {label} number[/yellow]").strip()
                if target.isdigit() and 1 <= int(target) <= count:
                    if view_mode == "paragraphs":
                        current_para = int(target) - 1
                    else:
                        current_table = int(target) - 1
                else:
                    console.print(f"[red]Invalid {label} number! Choose 1-{count}[/red]")
                    input("Press Enter to continue...")
            elif key == 't':
                view_mode = "tables" if view_mode == "paragraphs" else "paragraphs"
                if view_mode == "paragraphs":
                    current_para = 0
                else:
                    current_table = 0
            elif key == '/':
                search_term = Prompt.ask("[yellow]Enter search term[/yellow]").strip()
                if search_term:
                    search_results = []
                    for i, para in enumerate(paragraphs):
                        if search_term.lower() in para.text.lower():
                            search_results.append(("paragraph", i))
                    for i, table in enumerate(tables):
                        for row in table.rows:
                            for cell in row.cells:
                                if search_term.lower() in cell.text.lower():
                                    search_results.append(("table", i))
                                    break
                    
                    search_index = 0
                    if search_results:
                        result_type, result_index = search_results[0]
                        if result_type == "paragraph":
                            view_mode = "paragraphs"
                            current_para = result_index
                        else:
                            view_mode = "tables"
                            current_table = result_index
                    else:
                        console.print(f"[yellow]No results found for '{search_term}'[/yellow]")
                        input("Press Enter to continue...")
            elif key == 'n' and search_results:
                search_index = (search_index + 1) % len(search_results)
                result_type, result_index = search_results[search_index]
                if result_type == "paragraph":
                    view_mode = "paragraphs"
                    current_para = result_index
                else:
                    view_mode = "tables"
                    current_table = result_index
            elif key == 'p' and search_results:
                search_index = (search_index - 1) % len(search_results)
                result_type, result_index = search_results[search_index]
                if result_type == "paragraph":
                    view_mode = "paragraphs"
                    current_para = result_index
                else:
                    view_mode = "tables"
                    current_table = result_index
    
    except Exception as e:
        console.print(f"[red]Error reading DOCX: {e}[/red]")
        input("Press Enter to continue...")

def interactive_csv_viewer(path, filename, start_row=0):
    """Interactive CSV viewer with navigation and search"""
    reason = detect_binary_content(path)
    if reason:
        console.print(Panel(
            f"[bold red]'{filename}' doesn't look like a real CSV[/bold red]\n\n"
            f"{reason}.\n\n"
            "[dim]Re-export it as a plain-text CSV (or upload the original "
            "file with its correct extension) to view it here.[/dim]",
            expand=False,
        ))
        input("Press Enter to continue...")
        return

    try:
        with open(path, "r", newline="", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            rows = list(reader)

            if not rows:
                console.print("[yellow]CSV file is empty[/yellow]")
                input("Press Enter to continue...")
                return

            header = rows[0]
            data_rows = rows[1:] if len(rows) > 1 else []

            rows_per_page = 10
            current_row = 0
            if data_rows and start_row:
                current_row = (max(0, min(start_row - 1, len(data_rows) - 1)) // rows_per_page) * rows_per_page
            search_term = ""
            search_results = []
            search_index = 0

            while True:
                clear_screen()
                console.print(Panel(f"[bold cyan]📊 Interactive CSV Viewer[/bold cyan]", expand=False))
                console.print(f"[yellow]File:[/yellow] {filename}")
                console.print(f"[yellow]Total data rows:[/yellow] {len(data_rows)} (+1 header row)")
                if data_rows:
                    console.print(f"[yellow]Showing rows:[/yellow] {current_row + 1}-{min(current_row + rows_per_page, len(data_rows))}")

                if search_term:
                    console.print(f"[yellow]Search:[/yellow] '{search_term}' ({len(search_results)} results)")

                console.print()

                table = Table(show_header=True, header_style="bold magenta")
                table.add_column("Row", justify="right", width=6)
                for col_idx, col_name in enumerate(header):
                    table.add_column(col_name or f"Column {col_idx + 1}", overflow="fold")

                for i in range(current_row, min(current_row + rows_per_page, len(data_rows))):
                    row = data_rows[i]
                    is_hit = bool(search_term) and any(search_term.lower() in str(c).lower() for c in row)
                    row_label = f"[bold red]{i + 1}[/bold red]" if is_hit else str(i + 1)

                    cells = []
                    for col_idx in range(len(header)):
                        cell = row[col_idx] if col_idx < len(row) else ""
                        cell_text = str(cell)
                        if search_term and search_term.lower() in cell_text.lower():
                            cell_text = cell_text.replace(search_term, f"[bold red]{search_term}[/bold red]")
                        cells.append(cell_text)
                    table.add_row(row_label, *cells)

                console.print(table)

                console.print()
                print_nav_table([
                    ("↑/↓ k/j", "Previous / next page"),
                    ("Home/End g/G", "First / last page"),
                    (":", "Go to a specific row number"),
                    ("/", "Search in CSV"),
                    ("n / p", "Next / previous search result"),
                    ("q", "Quit"),
                ])

                key = get_key()

                if key == 'q':
                    break
                elif key == '\x1b[A' or key == 'k':  # Up arrow or k
                    current_row = max(0, current_row - rows_per_page)
                elif key == '\x1b[B' or key == 'j':  # Down arrow or j
                    current_row = min(max(0, len(data_rows) - rows_per_page), current_row + rows_per_page)
                elif key == '\x1b[H' or key == 'g':  # Home or g
                    current_row = 0
                elif key == '\x1b[F' or key == 'G':  # End or G
                    current_row = max(0, len(data_rows) - rows_per_page)
                elif key == ':':
                    target = Prompt.ask("[yellow]Go to row number[/yellow]").strip()
                    if target.isdigit() and 1 <= int(target) <= len(data_rows):
                        current_row = ((int(target) - 1) // rows_per_page) * rows_per_page
                    else:
                        console.print(f"[red]Invalid row number! Choose 1-{len(data_rows)}[/red]")
                        input("Press Enter to continue...")
                elif key == '/':
                    search_term = Prompt.ask("[yellow]Enter search term[/yellow]").strip()
                    if search_term:
                        search_results = []
                        for i, row in enumerate(data_rows):
                            for cell in row:
                                if search_term.lower() in str(cell).lower():
                                    search_results.append(i)
                                    break
                        search_index = 0
                        if search_results:
                            current_row = (search_results[0] // rows_per_page) * rows_per_page
                        else:
                            console.print(f"[yellow]No results found for '{search_term}'[/yellow]")
                            input("Press Enter to continue...")
                elif key == 'n' and search_results:
                    search_index = (search_index + 1) % len(search_results)
                    current_row = (search_results[search_index] // rows_per_page) * rows_per_page
                elif key == 'p' and search_results:
                    search_index = (search_index - 1) % len(search_results)
                    current_row = (search_results[search_index] // rows_per_page) * rows_per_page

    except Exception as e:
        console.print(f"[red]Error reading CSV: {e}[/red]")
        input("Press Enter to continue...")

def view_file_rich(user_folder, subject, filename, editable=False, jump_to=None):
    """Main file viewer function - handles ALL file types in CLI.
    jump_to, when given, is a 1-based line/row/page/paragraph number to open
    straight to (e.g. from a /search result) instead of the start of the file."""
    path = note_path(user_folder, subject, filename)
    
    if not os.path.exists(path):
        console.print(f"[red]File '{filename}' not found[/red]")
        return
    
    try:
        ext = filename.split(".")[-1].lower() if "." in filename else "txt"
        console.print(Panel(f"[bold cyan]Opening {filename} in CLI[/bold cyan]", expand=False))
        
        if editable:
            edit_file_with_reason(user_folder, subject, filename)
            return

        # Text-based files (interactive reader)
        if ext in TEXT_EXTENSIONS:
            interactive_text_reader(path, filename, start_line=(jump_to - 1) if jump_to else 0)
            
        # PDF files
        elif ext == "pdf":
            if not PyPDF2:
                console.print("[red]PyPDF2 not installed - showing file info instead[/red]")
                file_size = os.path.getsize(path)
                console.print(f"[cyan]PDF file: {filename}[/cyan]")
                console.print(f"[cyan]Size: {file_size:,} bytes[/cyan]")
                console.print()
                console.print("[yellow]To view PDF content:[/yellow]")
                console.print("[cyan]pip install PyPDF2[/cyan]")
                input("\nPress Enter to continue...")
                return
            
            # Interactive PDF viewer
            interactive_pdf_viewer(path, filename, start_page=(jump_to - 1) if jump_to else 0)

        # Word documents
        elif ext in ["doc", "docx"]:
            if not docx:
                console.print("[red]python-docx not installed - showing file info instead[/red]")
                file_size = os.path.getsize(path)
                console.print(f"[cyan]Word document: {filename}[/cyan]")
                console.print(f"[cyan]Size: {file_size:,} bytes[/cyan]")
                console.print()
                console.print("[yellow]To view Word document content:[/yellow]")
                console.print("[cyan]pip install python-docx lxml[/cyan]")
                console.print()
                console.print("[dim]Note: lxml is required for python-docx to work properly[/dim]")
                input("\nPress Enter to continue...")
                return
            
            # Validate document first
            console.print("[yellow]🔍 Validating document...[/yellow]")
            is_valid, message, has_text, has_tables = validate_word_document(path)
            
            if not is_valid:
                console.print(f"[red]❌ Document validation failed: {message}[/red]")
                console.print()
                
                # Provide specific error diagnosis
                if "corrupted or encrypted" in message:
                    console.print("[yellow]🔍 Diagnosis: Document structure issue[/yellow]")
                    console.print("This usually means:")
                    console.print("• Document is corrupted or damaged")
                    console.print("• Document is encrypted or password-protected")
                    console.print("• Document contains unsupported elements")
                    console.print("• Document was created with a very old/new version of Word")
                elif "lxml" in message.lower():
                    console.print("[yellow]🔍 Diagnosis: Missing XML processing library[/yellow]")
                    console.print("• lxml is required for python-docx to work properly")
                else:
                    console.print("[yellow]🔍 Diagnosis: General document reading error[/yellow]")
                    console.print("This might be due to:")
                    console.print("• Corrupted or encrypted document")
                    console.print("• Unsupported Word format")
                    console.print("• Missing dependencies")
                
                console.print()
                console.print("[cyan]💡 Solutions to try:[/cyan]")
                console.print("1. [yellow]Install/update lxml:[/yellow] pip install --upgrade lxml")
                console.print("2. [yellow]Try opening in Word and save as .docx again[/yellow]")
                console.print("3. [yellow]Convert to PDF and upload the PDF instead[/yellow]")
                console.print("4. [yellow]Check if document is password-protected[/yellow]")
                console.print("5. [yellow]Try uploading a different Word document[/yellow]")
                
                # Show file info as fallback
                try:
                    file_size = os.path.getsize(path)
                    console.print()
                    console.print(f"[cyan]File info:[/cyan] {filename}")
                    console.print(f"[cyan]Size:[/cyan] {file_size:,} bytes")
                    console.print(f"[cyan]Path:[/cyan] {path}")
                except:
                    pass
                
                input("\nPress Enter to continue...")
                return
            
            # Document is valid, use interactive viewer
            console.print(f"[green]✅ Document validated successfully[/green]")
            console.print(f"[cyan]Document has {len(doc.paragraphs)} paragraphs[/cyan]")
            
            if has_tables:
                console.print(f"[cyan]Document has {len(doc.tables)} tables[/cyan]")
            
            console.print()
            console.print("[yellow]Opening interactive DOCX viewer...[/yellow]")
            input("Press Enter to continue...")
            
            # Use interactive DOCX viewer
            interactive_docx_viewer(path, filename, start_para=(jump_to - 1) if jump_to else 0)

        # CSV files
        elif ext == "csv":
            interactive_csv_viewer(path, filename, start_row=jump_to or 0)

        # Image files
        elif ext in ["jpg", "jpeg", "png", "gif", "bmp", "tiff", "svg", "webp"]:
            view_image_info(path, filename, ext)

        # Archive files
        elif ext in ["zip", "rar", "7z", "tar", "gz", "bz2", "xz"]:
            view_archive_info(path, filename, ext)

        # Executable files
        elif ext in ["exe", "dll", "so", "dylib", "bin"]:
            console.print(f"[yellow]Executable file: {filename}[/yellow]")
            console.print("[yellow]Cannot execute or view binary executables in CLI.[/yellow]")
            input("Press Enter to continue...")

        # Audio/Video files
        elif ext in ["mp3", "wav", "flac", "mp4", "avi", "mkv", "mov", "wmv"]:
            file_size = os.path.getsize(path)
            console.print(f"[cyan]Media file: {filename}[/cyan]")
            console.print(f"[cyan]Format: {ext.upper()}[/cyan]")
            console.print(f"[cyan]Size: {file_size:,} bytes[/cyan]")
            console.print("[yellow]Media files cannot be played in CLI.[/yellow]")
            input("\nPress Enter to continue...")

        # Default: try to read as text, fallback to binary
        else:
            try:
                # First try to read as text
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read(1000)  # Read first 1000 chars
                    if content.strip():
                        console.print(f"[yellow]Unknown text format '{ext}' - showing as text:[/yellow]")
                        console.print(content)
                        if len(content) == 1000:
                            console.print("\n[yellow]... (showing first 1000 characters)[/yellow]")
                        input("\nPress Enter to continue...")
                    else:
                        raise UnicodeDecodeError("", b"", 0, 1, "empty content")
            except (UnicodeDecodeError, UnicodeError):
                # If text reading fails, show as binary
                console.print(f"[yellow]Unknown binary format '{ext}' - showing hex dump:[/yellow]")
                view_binary_file(path, filename, ext)

    except Exception as e:
        console.print(f"[red]Unexpected error opening file: {e}[/red]")
        console.print("[yellow]File may be corrupted or in an unsupported format[/yellow]")
        input("Press Enter to continue...")

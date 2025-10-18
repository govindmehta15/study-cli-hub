#!/usr/bin/env python3
"""
Document Repair Utility for CLI Study Hub
Helps fix common Word document issues
"""

import os
import sys
from rich.console import Console
from rich.panel import Panel
from rich.prompt import Prompt
from rich.table import Table

console = Console()

def check_document_issues(file_path):
    """Check for common document issues"""
    issues = []
    
    try:
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            issues.append("File is empty (0 bytes)")
        elif file_size < 100:
            issues.append("File is very small (might be corrupted)")
        
        # Check file extension
        if not file_path.lower().endswith(('.doc', '.docx')):
            issues.append("File doesn't have .doc or .docx extension")
        
        # Try to import python-docx
        try:
            import docx
        except ImportError:
            issues.append("python-docx library not installed")
            return issues
        
        # Try to open the document
        try:
            doc = docx.Document(file_path)
            
            # Check for content
            has_text = any(para.text.strip() for para in doc.paragraphs)
            has_tables = len(doc.tables) > 0
            
            if not has_text and not has_tables:
                issues.append("Document appears to be empty")
            elif not has_text and has_tables:
                issues.append("Document has tables but no text content")
            
            # Check for very long paragraphs (might indicate corruption)
            for para in doc.paragraphs:
                if len(para.text) > 10000:
                    issues.append("Document has unusually long paragraphs (possible corruption)")
                    break
                    
        except Exception as e:
            error_msg = str(e)
            if "Package not found" in error_msg:
                issues.append("Document structure is corrupted or encrypted")
            elif "lxml" in error_msg.lower():
                issues.append("lxml library not installed")
            else:
                issues.append(f"Document reading error: {error_msg}")
                
    except Exception as e:
        issues.append(f"File access error: {e}")
    
    return issues

def suggest_repairs(issues):
    """Suggest repair methods based on issues found"""
    suggestions = []
    
    for issue in issues:
        if "python-docx" in issue:
            suggestions.append("Install python-docx: pip install python-docx")
        elif "lxml" in issue:
            suggestions.append("Install lxml: pip install lxml")
        elif "empty" in issue.lower():
            suggestions.append("Document might be empty - check if it has content in Word")
        elif "corrupted" in issue.lower():
            suggestions.append("Try opening in Word and saving as new .docx file")
            suggestions.append("Convert to PDF and upload the PDF instead")
        elif "encrypted" in issue.lower():
            suggestions.append("Document might be password-protected - remove password in Word")
        elif "extension" in issue.lower():
            suggestions.append("Rename file to have .docx extension")
    
    return list(set(suggestions))  # Remove duplicates

def repair_document():
    """Main repair function"""
    console.print(Panel("[bold cyan]🔧 Document Repair Utility[/bold cyan]", expand=False))
    console.print()
    
    # Get file path
    file_path = Prompt.ask("[yellow]Enter path to Word document to repair[/yellow]").strip()
    
    if not os.path.exists(file_path):
        console.print(f"[red]File not found: {file_path}[/red]")
        return False
    
    console.print()
    console.print(f"[cyan]Analyzing: {os.path.basename(file_path)}[/cyan]")
    
    # Check for issues
    issues = check_document_issues(file_path)
    
    if not issues:
        console.print("[green]✅ No issues found! Document appears to be healthy.[/green]")
        return True
    
    # Display issues
    console.print()
    console.print(Panel("[bold red]Issues Found[/bold red]", expand=False))
    
    table = Table(title="Document Issues", show_header=True, header_style="bold red")
    table.add_column("Issue", width=50)
    table.add_column("Severity", width=15)
    
    for issue in issues:
        severity = "High" if any(word in issue.lower() for word in ["corrupted", "encrypted", "error"]) else "Medium"
        table.add_row(issue, severity)
    
    console.print(table)
    
    # Suggest repairs
    suggestions = suggest_repairs(issues)
    
    if suggestions:
        console.print()
        console.print(Panel("[bold yellow]Suggested Repairs[/bold yellow]", expand=False))
        
        for i, suggestion in enumerate(suggestions, 1):
            console.print(f"{i}. [cyan]{suggestion}[/cyan]")
    
    console.print()
    console.print("[yellow]💡 Additional Tips:[/yellow]")
    console.print("• Try opening the document in Microsoft Word first")
    console.print("• Save as a new .docx file in Word")
    console.print("• Convert to PDF if Word document issues persist")
    console.print("• Check if document is password-protected")
    console.print("• Ensure document was not created with very old Word version")
    
    return False

def main():
    """Main function"""
    console.print("CLI Study Hub - Document Repair Utility")
    console.print("=" * 50)
    console.print()
    
    try:
        repair_document()
    except KeyboardInterrupt:
        console.print("\n[yellow]Repair cancelled by user[/yellow]")
    except Exception as e:
        console.print(f"[red]Error: {e}[/red]")
    
    console.print()
    input("Press Enter to continue...")

if __name__ == "__main__":
    main()
